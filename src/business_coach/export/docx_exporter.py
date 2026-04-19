"""DOCX generation for mPABS."""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

def _safe_add_heading(doc: Document, text: str, level: int = 1) -> None:
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        size_map = {1: Pt(18), 2: Pt(16), 3: Pt(14), 4: Pt(12)}
        run.font.size = size_map.get(level, Pt(14))

def _add_markdown_content(doc: Document, text: str) -> None:
    # Simplified markdown to docx for business coach
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            _safe_add_heading(doc, line[3:], 2)
        elif line.startswith("### "):
            _safe_add_heading(doc, line[4:], 3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

class DOCXExporter:
    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_canvas(self, topic_name: str, canvas_elements: list) -> Path:
        doc = Document()
        _safe_add_heading(doc, f"Business Model Canvas: {topic_name}", level=1)
        
        for el in canvas_elements:
            _safe_add_heading(doc, el.element_name, level=2)
            _add_markdown_content(doc, el.content)
            
        out_path = self.output_dir / f"{topic_name.replace(' ', '_')}_Canvas.docx"
        doc.save(str(out_path))
        return out_path

    def export_voices(self, topic_name: str, voices: list) -> Path:
        doc = Document()
        _safe_add_heading(doc, f"Target Personas: {topic_name}", level=1)
        
        for v in voices:
            _safe_add_heading(doc, v.name, level=2)
            p = doc.add_paragraph()
            p.add_run(f"Style: {v.communication_style}").italic = True
            _add_markdown_content(doc, v.description)
            
        out_path = self.output_dir / f"{topic_name.replace(' ', '_')}_Voices.docx"
        doc.save(str(out_path))
        return out_path

    def export_plan(self, topic_name: str, plan_sections: list) -> Path:
        doc = Document()
        _safe_add_heading(doc, f"Business Plan: {topic_name}", level=1)
        
        for sec in plan_sections:
            _safe_add_heading(doc, sec.section_name, level=2)
            _add_markdown_content(doc, sec.content)
            
        out_path = self.output_dir / f"{topic_name.replace(' ', '_')}_BusinessPlan.docx"
        doc.save(str(out_path))
        return out_path
